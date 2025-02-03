#!/usr/bin/env python3
"""
MIT License

Copyright (c) 2023 Darkspace Software & Security

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is furnished
to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.

Darkspace Software & Security  
Author: Michael Blenkinsop  
Email: mickyblenk@gmail.com  
Tel: +447710669684

Use Cases & Rationale:
  Traditional vulnerability tools (e.g., nmap, nikto, etc.) are frequently blocked by administrators.
  This tool uses alternative methods that are often not blocked—making it a must-use solution for vulnerability
  assessment in restricted environments. It also provides easy-to-read reports, patching recommendations with useful
  links, and the ability to scan specific ports.
"""

###############################################################################
# Automated Dependency Installer
###############################################################################
import subprocess
import sys

def install_missing_packages():
    # List of packages to check/install; for python-docx, we import as "docx"
    packages = [
        ("PyQt5", "PyQt5"),
        ("requests", "requests"),
        ("matplotlib", "matplotlib"),
        ("python-docx", "docx"),
        ("xlwt", "xlwt")
    ]
    for package_name, import_name in packages:
        try:
            __import__(import_name)
        except ImportError:
            print(f"Package {package_name} not found. Installing...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

install_missing_packages()

###############################################################################
# Now import all dependencies
###############################################################################
import re
import requests
import socket
import datetime
import time
from PyQt5 import QtCore, QtWidgets, QtGui
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from docx import Document
import xlwt

###############################################################################
# Global API configuration
###############################################################################
API_CONFIG = {
    "NVD": "",
    "CIRCL": "",
    "Vulners": "",
    "default_source": "CIRCL"  # Options: "NVD", "CIRCL", "Vulners"
}

###############################################################################
# Helper functions for scanning
###############################################################################
def get_installed_programs():
    programs = {}
    try:
        output = subprocess.check_output(
            ["wmic", "product", "get", "Name,Version"],
            universal_newlines=True
        )
        lines = output.strip().splitlines()
        if len(lines) > 1:
            for line in lines[1:]:
                if line.strip():
                    parts = re.split(r'\s{2,}', line.strip())
                    if len(parts) >= 2:
                        name, version = parts[0], parts[1]
                        programs[name] = version
                    elif len(parts) == 1:
                        programs[parts[0]] = ""
    except Exception as e:
        print("Error retrieving installed programs, using sample data:", e)
        programs = {
            "Mozilla Firefox": "102.0",
            "Google Chrome": "105.0",
            "Notepad++": "8.4.9"
        }
    return programs

###############################################################################
# Vulnerability query functions for multiple sources
###############################################################################
def query_circl_vulns(keyword):
    url = f"https://cve.circl.lu/api/search/{keyword}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            vulns = []
            for item in data.get("results", []):
                cve_id = item.get("id", "N/A")
                description = item.get("summary", "No description available.")
                severity = "UNKNOWN"
                cvss = item.get("cvss", None)
                if cvss is not None:
                    try:
                        cvss = float(cvss)
                        if cvss >= 9:
                            severity = "CRITICAL"
                        elif cvss >= 7:
                            severity = "HIGH"
                        elif cvss >= 4:
                            severity = "MEDIUM"
                        else:
                            severity = "LOW"
                    except:
                        severity = "UNKNOWN"
                vulns.append({
                    "cve_id": cve_id,
                    "description": description,
                    "severity": severity
                })
            return vulns
        else:
            print(f"CIRCL API response status: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error querying CIRCL API for keyword '{keyword}':", e)
        return []

def query_nvd_vulns(keyword):
    url = "https://services.nvd.nist.gov/rest/json/cves/1.0"
    params = {
        "keyword": keyword,
        "apiKey": API_CONFIG["NVD"],
        "resultsPerPage": 5
    }
    try:
        response = requests.get(url, params=params, timeout=10)
        if response.status_code == 200:
            data = response.json()
            vulns = []
            for item in data.get("result", {}).get("CVE_Items", []):
                cve_id = item.get("cve", {}).get("CVE_data_meta", {}).get("ID", "N/A")
                description_data = item.get("cve", {}).get("description", {}).get("description_data", [])
                description = description_data[0]["value"] if description_data else "No description available."
                severity = "UNKNOWN"
                impact = item.get("impact", {})
                if "baseMetricV3" in impact:
                    score = impact["baseMetricV3"]["cvssV3"]["baseScore"]
                    if score >= 9:
                        severity = "CRITICAL"
                    elif score >= 7:
                        severity = "HIGH"
                    elif score >= 4:
                        severity = "MEDIUM"
                    else:
                        severity = "LOW"
                elif "baseMetricV2" in impact:
                    sev = impact["baseMetricV2"]["severity"]
                    severity = sev.upper()
                vulns.append({
                    "cve_id": cve_id,
                    "description": description,
                    "severity": severity
                })
            return vulns
        else:
            print(f"NVD API response status: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error querying NVD API for keyword '{keyword}':", e)
        return []

def query_vulners_vulns(keyword):
    if not API_CONFIG["Vulners"]:
        print("No Vulners API key provided")
        return []
    url = "https://vulners.com/api/v3/search/lucene/"
    payload = {
        "query": keyword,
        "size": 5,
        "apiKey": API_CONFIG["Vulners"]
    }
    try:
        response = requests.post(url, json=payload, timeout=10)
        if response.status_code == 200:
            data = response.json()
            vulns = []
            for item in data.get("result", {}).get("documents", []):
                cve_id = item.get("id", "N/A")
                description = item.get("description", "No description available.")
                cvss = item.get("cvss", "UNKNOWN")
                severity = "UNKNOWN"
                try:
                    cvss = float(cvss)
                    if cvss >= 9:
                        severity = "CRITICAL"
                    elif cvss >= 7:
                        severity = "HIGH"
                    elif cvss >= 4:
                        severity = "MEDIUM"
                    else:
                        severity = "LOW"
                except:
                    severity = "UNKNOWN"
                vulns.append({
                    "cve_id": cve_id,
                    "description": description,
                    "severity": severity
                })
            return vulns
        else:
            print(f"Vulners API response status: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error querying Vulners API for keyword '{keyword}':", e)
        return []

def query_vulns(keyword):
    source = API_CONFIG.get("default_source", "CIRCL")
    if source == "NVD":
        return query_nvd_vulns(keyword)
    elif source == "Vulners":
        return query_vulners_vulns(keyword)
    else:
        return query_circl_vulns(keyword)

def aggregate_severity(scan_results):
    severity_counts = {}
    for prog, info in scan_results.items():
        for vuln in info.get("vulnerabilities", []):
            sev = vuln.get("severity", "UNKNOWN")
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
    return severity_counts

###############################################################################
# Port scanning functions (with option for specific ports)
###############################################################################
def scan_domain(domain, ports=None):
    vulns = []
    try:
        if ports is None:
            ports = [80, 443, 8080]
        else:
            ports = [int(p.strip()) for p in ports.split(",") if p.strip().isdigit()]
        open_ports = []
        for port in ports:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(1)
            result = sock.connect_ex((domain, port))
            if result == 0:
                open_ports.append(port)
            sock.close()
        if open_ports:
            vulns.append({"type": "open_ports", "details": f"Open ports: {', '.join(map(str, open_ports))}"})
        try:
            response = requests.get(f"http://{domain}", timeout=5)
            headers = response.headers
            if "Server" in headers:
                vulns.append({"type": "server_header", "details": f"Server: {headers['Server']}"})
            if "X-Frame-Options" not in headers:
                vulns.append({"type": "missing_header", "details": "X-Frame-Options header missing (clickjacking vulnerability)"})
            if "Content-Security-Policy" not in headers:
                vulns.append({"type": "missing_header", "details": "Content-Security-Policy header missing (XSS protection)"})
        except requests.RequestException as e:
            vulns.append({"type": "http_error", "details": f"HTTP request failed: {e}"})
        try:
            response = requests.get(f"http://{domain}/", timeout=5)
            if "Index of /" in response.text:
                vulns.append({"type": "directory_listing", "details": "Directory listing enabled"})
        except requests.RequestException as e:
            vulns.append({"type": "http_error", "details": f"HTTP request failed: {e}"})
    except Exception as e:
        vulns.append({"type": "error", "details": f"Error scanning domain: {e}"})
    return vulns

def detailed_scan_domain(domain, mode="Normal", ports=None):
    open_ports = []
    if mode == "Stealthy":
        timeout = 1.0
    elif mode == "Aggressive":
        timeout = 0.25
    else:
        timeout = 0.5
    if ports is None:
        ports = range(1, 1025)
    else:
        ports = [int(p.strip()) for p in ports.split(",") if p.strip().isdigit()]
    for port in ports:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(timeout)
            result = sock.connect_ex((domain, port))
            if result == 0:
                open_ports.append(port)
            sock.close()
        except Exception:
            continue
    return {"open_ports": open_ports}

###############################################################################
# Detailed Vulnerability Info Functions
###############################################################################
def get_detailed_vuln_info_circl(cve_id):
    url = f"https://cve.circl.lu/api/cve/{cve_id}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            published_date = data.get("Published", "N/A")
            last_modified_date = data.get("Modified", "N/A")
            cvss = data.get("cvss", "N/A")
            references = data.get("references", [])
            return {
                "published_date": published_date,
                "last_modified_date": last_modified_date,
                "cvss_v3_score": cvss,
                "cvss_v2_score": "N/A",
                "references": references
            }
        else:
            print(f"CIRCL detailed API response status: {response.status_code}")
            return None
    except Exception as e:
        print(f"Error fetching detailed vulnerability info for {cve_id}: {e}")
        return None

def get_detailed_vuln_info_nvd(cve_id):
    url = "https://services.nvd.nist.gov/rest/json/cves/1.0"
    params = {
        "cveId": cve_id,
        "apiKey": API_CONFIG["NVD"]
    }
    try:
        response = requests.get(url, params=params, timeout=10)
        if response.status_code == 200:
            data = response.json()
            items = data.get("result", {}).get("CVE_Items", [])
            if items:
                item = items[0]
                published_date = item.get("publishedDate", "N/A")
                last_modified_date = item.get("lastModifiedDate", "N/A")
                cvss_v3 = "N/A"
                if "baseMetricV3" in item.get("impact", {}):
                    cvss_v3 = item["impact"]["baseMetricV3"]["cvssV3"]["baseScore"]
                return {
                    "published_date": published_date,
                    "last_modified_date": last_modified_date,
                    "cvss_v3_score": cvss_v3,
                    "cvss_v2_score": "N/A",
                    "references": []
                }
            else:
                return None
        else:
            print(f"NVD detailed API response status: {response.status_code}")
            return None
    except Exception as e:
        print(f"Error fetching detailed vulnerability info for {cve_id}: {e}")
        return None

def get_detailed_vuln_info_vulners(cve_id):
    if not API_CONFIG["Vulners"]:
        print("No Vulners API key provided")
        return None
    url = "https://vulners.com/api/v3/search/lucene/"
    payload = {
        "query": f"id:{cve_id}",
        "size": 1,
        "apiKey": API_CONFIG["Vulners"]
    }
    try:
        response = requests.post(url, json=payload, timeout=10)
        if response.status_code == 200:
            data = response.json()
            docs = data.get("result", {}).get("documents", [])
            if docs:
                doc = docs[0]
                published_date = doc.get("published", "N/A")
                cvss = doc.get("cvss", "N/A")
                references = doc.get("references", [])
                return {
                    "published_date": published_date,
                    "last_modified_date": "N/A",
                    "cvss_v3_score": cvss,
                    "cvss_v2_score": "N/A",
                    "references": references
                }
            else:
                return None
        else:
            print(f"Vulners detailed API response status: {response.status_code}")
            return None
    except Exception as e:
        print(f"Error fetching detailed vulnerability info for {cve_id}: {e}")
        return None

def get_detailed_vuln_info(cve_id):
    source = API_CONFIG.get("default_source", "CIRCL")
    if source == "NVD":
        return get_detailed_vuln_info_nvd(cve_id)
    elif source == "Vulners":
        return get_detailed_vuln_info_vulners(cve_id)
    else:
        return get_detailed_vuln_info_circl(cve_id)

###############################################################################
# Worker Threads
###############################################################################
class ScanWorker(QtCore.QThread):
    progress = QtCore.pyqtSignal(str, str, list)
    finished = QtCore.pyqtSignal(dict)
    def __init__(self, mode="Normal", parent=None):
        super().__init__(parent)
        self.mode = mode
    def run(self):
        results = {}
        programs = get_installed_programs()
        for prog, version in programs.items():
            vulns = query_vulns(prog)
            results[prog] = {"version": version, "vulnerabilities": vulns}
            self.progress.emit(prog, version, vulns)
            if self.mode == "Stealthy":
                time.sleep(2)
            elif self.mode == "Normal":
                time.sleep(0.5)
            elif self.mode == "Aggressive":
                time.sleep(0)
        self.finished.emit(results)

class DomainScanWorker(QtCore.QThread):
    finished = QtCore.pyqtSignal(dict)
    def __init__(self, domain, ports=None, parent=None):
        super().__init__(parent)
        self.domain = domain
        self.ports = ports  # Optional comma-separated string of ports
    def run(self):
        results = {self.domain: scan_domain(self.domain, ports=self.ports)}
        self.finished.emit(results)

class DetailedDomainScanWorker(QtCore.QThread):
    finished = QtCore.pyqtSignal(dict)
    def __init__(self, domain, mode="Normal", ports=None, parent=None):
        super().__init__(parent)
        self.domain = domain
        self.mode = mode
        self.ports = ports  # Optional comma-separated string of ports
    def run(self):
        results = {self.domain: detailed_scan_domain(self.domain, self.mode, ports=self.ports)}
        self.finished.emit(results)

###############################################################################
# Patching Recommendations Dialog
###############################################################################
class PatchingDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Patching Recommendations")
        self.resize(600, 400)
        layout = QtWidgets.QVBoxLayout(self)
        recommendations = (
            "<h2>Patching Recommendations</h2>"
            "<p>This tool provides vulnerability data that can help you prioritize your patching efforts. "
            "For detailed patching advice and security best practices, consider visiting these resources:</p>"
            "<ul>"
            "<li><a href='https://www.cisa.gov/patching'>CISA Patching Guidance</a></li>"
            "<li><a href='https://www.us-cert.gov/ncas/tips/ST04-006'>US-CERT Vulnerability Tips</a></li>"
            "<li><a href='https://www.microsoft.com/en-us/security/portal/mmpc/shared/vulnerability.aspx'>Microsoft Security Updates</a></li>"
            "<li><a href='https://www.cisco.com/c/en/us/support/security/patch-advisory.html'>Cisco Patch Advisories</a></li>"
            "</ul>"
            "<p>Use these links to access developer tools and guidelines to remediate vulnerabilities effectively.</p>"
        )
        label = QtWidgets.QLabel(recommendations)
        label.setOpenExternalLinks(True)
        label.setWordWrap(True)
        layout.addWidget(label)
        close_btn = QtWidgets.QPushButton("Close")
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn, alignment=QtCore.Qt.AlignRight)

###############################################################################
# Settings Dialog for API Keys and Default Source
###############################################################################
class SettingsDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("API Settings")
        self.resize(500, 300)
        layout = QtWidgets.QFormLayout(self)
        self.nvd_key = QtWidgets.QLineEdit()
        self.nvd_key.setText(API_CONFIG.get("NVD", ""))
        layout.addRow("NVD API Key:", self.nvd_key)
        self.nvd_link = QtWidgets.QLabel('<a href="https://nvd.nist.gov/developers">Get NVD API Key</a>')
        self.nvd_link.setOpenExternalLinks(True)
        layout.addRow("", self.nvd_link)
        self.circl_key = QtWidgets.QLineEdit()
        self.circl_key.setText(API_CONFIG.get("CIRCL", ""))
        layout.addRow("CIRCL API Key:", self.circl_key)
        self.circl_link = QtWidgets.QLabel('<a href="https://cve.circl.lu/">CIRCL CVE API Info</a>')
        self.circl_link.setOpenExternalLinks(True)
        layout.addRow("", self.circl_link)
        self.vulners_key = QtWidgets.QLineEdit()
        self.vulners_key.setText(API_CONFIG.get("Vulners", ""))
        layout.addRow("Vulners API Key:", self.vulners_key)
        self.vulners_link = QtWidgets.QLabel('<a href="https://vulners.com/documentation">Get Vulners API Key</a>')
        self.vulners_link.setOpenExternalLinks(True)
        layout.addRow("", self.vulners_link)
        self.source_combo = QtWidgets.QComboBox()
        self.source_combo.addItems(["NVD", "CIRCL", "Vulners"])
        self.source_combo.setCurrentText(API_CONFIG.get("default_source", "CIRCL"))
        layout.addRow("Default Source:", self.source_combo)
        button_box = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addRow(button_box)
    def get_settings(self):
        return {
            "NVD": self.nvd_key.text().strip(),
            "CIRCL": self.circl_key.text().strip(),
            "Vulners": self.vulners_key.text().strip(),
            "default_source": self.source_combo.currentText()
        }

###############################################################################
# Main Window with Modern GUI
###############################################################################
class VulnScannerWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Modern Vulnerability Scanner")
        self.resize(900, 600)
        self.scan_results = {}
        self.domain_scan_results = {}
        self.dark_mode = False
        self.worker = None
        self.domain_worker = None
        self.detailed_domain_worker = None

        # Set blue style for all buttons
        self.setStyleSheet("QPushButton { background-color: blue; color: white; }")

        # Create menu bar with About and Patching Recommendations
        menu_bar = self.menuBar()
        help_menu = menu_bar.addMenu("Help")
        about_action = QtWidgets.QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        patching_action = QtWidgets.QAction("Patching Recommendations", self)
        patching_action.triggered.connect(self.show_patching_recommendations)
        help_menu.addAction(patching_action)

        central_widget = QtWidgets.QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QtWidgets.QVBoxLayout(central_widget)

        # Create a tab widget for scan results
        self.tabs = QtWidgets.QTabWidget()
        main_layout.addWidget(self.tabs)

        self.local_scan_tab = QtWidgets.QWidget()
        self.tabs.addTab(self.local_scan_tab, "Local Scan")
        local_layout = QtWidgets.QVBoxLayout(self.local_scan_tab)
        self.local_tree = QtWidgets.QTreeWidget()
        self.local_tree.setColumnCount(3)
        self.local_tree.setHeaderLabels(["Program / CVE ID", "Version", "Severity"])
        self.local_tree.itemDoubleClicked.connect(self.on_item_double_clicked)
        local_layout.addWidget(self.local_tree)

        self.domain_scan_tab = QtWidgets.QWidget()
        self.tabs.addTab(self.domain_scan_tab, "Domain Scan")
        domain_layout = QtWidgets.QVBoxLayout(self.domain_scan_tab)
        domain_input_layout = QtWidgets.QHBoxLayout()
        self.domain_input = QtWidgets.QLineEdit()
        self.domain_input.setPlaceholderText("Enter domain to scan")
        domain_input_layout.addWidget(self.domain_input)
        self.port_input = QtWidgets.QLineEdit()
        self.port_input.setPlaceholderText("Optional: Enter comma-separated ports (e.g., 22,80,443)")
        domain_input_layout.addWidget(self.port_input)
        self.domain_scan_button = QtWidgets.QPushButton("Scan Domain")
        self.domain_scan_button.clicked.connect(self.start_domain_scan)
        domain_input_layout.addWidget(self.domain_scan_button)
        self.detailed_domain_scan_button = QtWidgets.QPushButton("Detailed Domain Scan")
        self.detailed_domain_scan_button.clicked.connect(self.start_detailed_domain_scan)
        domain_input_layout.addWidget(self.detailed_domain_scan_button)
        domain_layout.addLayout(domain_input_layout)
        self.domain_scan_output = QtWidgets.QTextEdit()
        self.domain_scan_output.setReadOnly(True)
        domain_layout.addWidget(self.domain_scan_output)

        # Scan Options Group Box
        scan_options_group = QtWidgets.QGroupBox("Scan Options")
        options_layout = QtWidgets.QHBoxLayout(scan_options_group)
        self.normal_radio = QtWidgets.QRadioButton("Normal")
        self.normal_radio.setChecked(True)
        self.stealthy_radio = QtWidgets.QRadioButton("Stealthy")
        self.aggressive_radio = QtWidgets.QRadioButton("Aggressive")
        options_layout.addWidget(self.normal_radio)
        options_layout.addWidget(self.stealthy_radio)
        options_layout.addWidget(self.aggressive_radio)
        main_layout.addWidget(scan_options_group)

        # Create toolbar with actions
        toolbar = QtWidgets.QToolBar("Main Toolbar")
        self.addToolBar(toolbar)
        scan_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("system-run"), "Scan System", self)
        scan_action.triggered.connect(self.start_scan)
        toolbar.addAction(scan_action)
        stop_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("process-stop"), "Stop Scan", self)
        stop_action.triggered.connect(self.stop_scans)
        toolbar.addAction(stop_action)
        report_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("document-save"), "Download Report", self)
        report_action.triggered.connect(self.download_report)
        toolbar.addAction(report_action)
        chart_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("view-statistics"), "Show Chart", self)
        chart_action.triggered.connect(self.show_chart)
        toolbar.addAction(chart_action)
        detailed_info_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("help-about"), "Detailed Vulnerability Info", self)
        detailed_info_action.triggered.connect(self.show_detailed_info)
        toolbar.addAction(detailed_info_action)
        settings_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("preferences-system"), "Settings", self)
        settings_action.triggered.connect(self.open_settings)
        toolbar.addAction(settings_action)
        toggle_theme_action = QtWidgets.QAction(QtGui.QIcon.fromTheme("preferences-desktop-theme"), "Toggle Dark Mode", self)
        toggle_theme_action.triggered.connect(self.toggle_dark_mode)
        toolbar.addAction(toggle_theme_action)

        self.status = self.statusBar()
        self.status.showMessage("Ready")

    def show_about(self):
        about_text = (
            "<b>Darkspace Software & Security</b><br>"
            "Author: Michael Blenkinsop<br>"
            "Email: mickyblenk@gmail.com<br>"
            "Tel: +447710669684<br><br>"
            "<b>Why This Tool?</b><br>"
            "Traditional tools like nmap, nikto, and others are frequently blocked by administrators, "
            "rendering them ineffective. This scanner uses alternative methods that are often not blocked, making "
            "it a must-use solution for vulnerability assessment in restricted environments."
        )
        QtWidgets.QMessageBox.about(self, "About", about_text)

    def show_patching_recommendations(self):
        dialog = PatchingDialog(self)
        dialog.exec_()

    def open_settings(self):
        dialog = SettingsDialog(self)
        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            new_settings = dialog.get_settings()
            API_CONFIG.update(new_settings)
            self.status.showMessage("API settings updated", 5000)

    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            dark_style = (
                "QWidget { background-color: #2b2b2b; color: #d3d3d3; }"
                "QTreeWidget { background-color: #3c3c3c; }"
                "QHeaderView::section { background-color: #3c3c3c; color: #d3d3d3; }"
                "QToolBar { background-color: #3c3c3c; }"
                "QMenuBar { background-color: #2b2b2b; color: #d3d3d3; }"
                "QMenu { background-color: #2b2b2b; color: #d3d3d3; }"
                "QPushButton { background-color: blue; color: white; border: 1px solid #666; }"
            )
            self.setStyleSheet(dark_style)
        else:
            self.setStyleSheet("QPushButton { background-color: blue; color: white; }")

    def get_scan_mode(self):
        if self.stealthy_radio.isChecked():
            return "Stealthy"
        elif self.aggressive_radio.isChecked():
            return "Aggressive"
        else:
            return "Normal"

    def start_scan(self):
        self.local_tree.clear()
        self.scan_results = {}
        self.status.showMessage("Scanning local system...")
        mode = self.get_scan_mode()
        self.worker = ScanWorker(mode)
        self.worker.progress.connect(self.update_local_tree)
        self.worker.finished.connect(self.scan_finished)
        self.worker.start()

    def start_domain_scan(self):
        domain = self.domain_input.text().strip()
        if not domain:
            QtWidgets.QMessageBox.warning(self, "Input Error", "Please enter a domain to scan.")
            return
        self.domain_scan_output.clear()
        self.domain_scan_results = {}
        self.status.showMessage(f"Scanning domain {domain}...")
        ports = self.port_input.text().strip()  # May be empty
        self.domain_worker = DomainScanWorker(domain, ports=ports if ports else None)
        self.domain_worker.finished.connect(self.domain_scan_finished)
        self.domain_worker.start()

    def start_detailed_domain_scan(self):
        domain = self.domain_input.text().strip()
        if not domain:
            QtWidgets.QMessageBox.warning(self, "Input Error", "Please enter a domain for detailed scan.")
            return
        self.domain_scan_output.clear()
        self.domain_scan_results = {}
        self.status.showMessage(f"Performing detailed scan on {domain}...")
        mode = self.get_scan_mode()
        ports = self.port_input.text().strip()  # May be empty
        self.detailed_domain_worker = DetailedDomainScanWorker(domain, mode, ports=ports if ports else None)
        self.detailed_domain_worker.finished.connect(self.detailed_domain_scan_finished)
        self.detailed_domain_worker.start()

    def stop_scans(self):
        if self.worker and self.worker.isRunning():
            self.worker.terminate()
            self.worker.wait()
        if self.domain_worker and self.domain_worker.isRunning():
            self.domain_worker.terminate()
            self.domain_worker.wait()
        if self.detailed_domain_worker and self.detailed_domain_worker.isRunning():
            self.detailed_domain_worker.terminate()
            self.detailed_domain_worker.wait()
        self.status.showMessage("Scan stopped")

    def update_local_tree(self, prog, version, vulns):
        parent = QtWidgets.QTreeWidgetItem([prog, version, ""])
        if vulns:
            for vuln in vulns:
                child = QtWidgets.QTreeWidgetItem([vuln.get("cve_id", "N/A"), "", vuln.get("severity", "UNKNOWN")])
                parent.addChild(child)
        self.local_tree.addTopLevelItem(parent)

    def scan_finished(self, results):
        self.scan_results = results
        self.status.showMessage("Local scan complete")

    def domain_scan_finished(self, results):
        self.domain_scan_results = results
        self.status.showMessage("Domain scan complete")
        domain = list(results.keys())[0]
        vulns = results[domain]
        self.domain_scan_output.append(f"Quick scan results for {domain}:\n")
        for vuln in vulns:
            self.domain_scan_output.append(f"Type: {vuln['type']}\n")
            self.domain_scan_output.append(f"Details: {vuln['details']}\n")
            self.domain_scan_output.append("\n")

    def detailed_domain_scan_finished(self, results):
        self.domain_scan_results = results
        self.status.showMessage("Detailed domain scan complete")
        domain = list(results.keys())[0]
        details = results[domain]
        open_ports = details.get("open_ports", [])
        self.domain_scan_output.append(f"Detailed scan results for {domain}:\n")
        if open_ports:
            self.domain_scan_output.append(f"Open ports: {', '.join(map(str, open_ports))}\n")
        else:
            self.domain_scan_output.append("No open ports found.\n")

    def download_report(self):
        if not self.scan_results and not self.domain_scan_results:
            QtWidgets.QMessageBox.information(self, "No Data", "No scan results available.")
            return

        report_format, ok = QtWidgets.QInputDialog.getItem(
            self, "Select Report Format", "Choose format:", ["docx", "xls", "txt", "html"], 2, False
        )
        if not ok:
            return

        report_lines = []
        report_lines.append("Vulnerability Scan Report")
        report_lines.append("Date: " + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        report_lines.append("")
        if self.scan_results:
            report_lines.append("Local Scan Results:")
            for prog, info in self.scan_results.items():
                version = info.get("version", "")
                vulns = info.get("vulnerabilities", [])
                report_lines.append(f"Program: {prog} (Version: {version})")
                if vulns:
                    for vuln in vulns:
                        report_lines.append(f"  - CVE ID: {vuln.get('cve_id', 'N/A')}")
                        report_lines.append(f"    Severity: {vuln.get('severity', 'UNKNOWN')}")
                        report_lines.append(f"    Description: {vuln.get('description', 'No description available.')}")
                else:
                    report_lines.append("  No vulnerabilities found.")
                report_lines.append("")
        if self.domain_scan_results:
            report_lines.append("Domain Scan Results:")
            for domain, result in self.domain_scan_results.items():
                report_lines.append(f"Domain: {domain}")
                if isinstance(result, list):
                    for vuln in result:
                        report_lines.append(f"  - Type: {vuln.get('type', 'N/A')}")
                        report_lines.append(f"    Details: {vuln.get('details', '')}")
                elif isinstance(result, dict):
                    open_ports = result.get("open_ports", [])
                    report_lines.append(f"  Open ports: {', '.join(map(str, open_ports)) if open_ports else 'None'}")
                report_lines.append("")
        report_text = "\n".join(report_lines)

        if report_format == "txt":
            file_filter = "Text Files (*.txt)"
            extension = ".txt"
        elif report_format == "html":
            file_filter = "HTML Files (*.html)"
            extension = ".html"
        elif report_format == "docx":
            file_filter = "Word Documents (*.docx)"
            extension = ".docx"
        elif report_format == "xls":
            file_filter = "Excel Files (*.xls)"
            extension = ".xls"
        else:
            file_filter = "Text Files (*.txt)"
            extension = ".txt"

        filename, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Report", "", file_filter)
        if not filename:
            return
        if not filename.lower().endswith(extension):
            filename += extension

        try:
            if report_format == "txt":
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(report_text)
            elif report_format == "html":
                html_content = "<html><body><pre>" + report_text + "</pre></body></html>"
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(html_content)
            elif report_format == "docx":
                if Document is None:
                    QtWidgets.QMessageBox.warning(self, "Error", "python-docx module not installed. Install with 'pip install python-docx'.")
                    return
                document = Document()
                document.add_heading("Vulnerability Scan Report", 0)
                document.add_paragraph("Date: " + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                for line in report_lines:
                    document.add_paragraph(line)
                document.save(filename)
            elif report_format == "xls":
                if xlwt is None:
                    QtWidgets.QMessageBox.warning(self, "Error", "xlwt module not installed. Install with 'pip install xlwt'.")
                    return
                workbook = xlwt.Workbook()
                sheet = workbook.add_sheet("Report")
                for row, line in enumerate(report_lines):
                    sheet.write(row, 0, line)
                workbook.save(filename)
            QtWidgets.QMessageBox.information(self, "Success", f"Report saved to {filename}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Failed to save report: {e}")

    def show_chart(self):
        if not self.scan_results:
            QtWidgets.QMessageBox.information(self, "No Data", "No local scan results available for charting.")
            return
        severity_counts = aggregate_severity(self.scan_results)
        if not severity_counts:
            QtWidgets.QMessageBox.information(self, "No Data", "No vulnerabilities found in local scan results.")
            return
        chart_win = ChartWindow(severity_counts, self)
        chart_win.exec_()

    def on_item_double_clicked(self, item, column):
        parent = item.parent()
        if parent:
            prog = parent.text(0)
            cve_id = item.text(0)
            vulns = self.scan_results.get(prog, {}).get("vulnerabilities", [])
            vuln_detail = next((v for v in vulns if v.get("cve_id") == cve_id), None)
            if vuln_detail:
                exploit_win = ExploitWindow(prog, vuln_detail, self)
                exploit_win.exec_()

    def show_detailed_info(self):
        selected_items = self.local_tree.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.information(self, "No Selection", "Please select a vulnerability (child node) first.")
            return
        item = selected_items[0]
        parent = item.parent()
        if not parent:
            QtWidgets.QMessageBox.information(self, "Selection Error", "Please select a vulnerability (child node), not a program.")
            return
        cve_id = item.text(0)
        self.status.showMessage(f"Fetching detailed info for {cve_id}...")
        detailed_info = get_detailed_vuln_info(cve_id)
        self.status.showMessage("Detailed info fetched.")
        detailed_win = DetailedVulnWindow(cve_id, detailed_info, self)
        detailed_win.exec_()

def main():
    app = QtWidgets.QApplication(sys.argv)
    window = VulnScannerWindow()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
